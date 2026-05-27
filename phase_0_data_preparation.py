"""
Phase 0: Data Preparation (Pre-processing)
Converts raw .docx files into clean, structured JSONL format.

Output: One JSONL file per source document with structured elements.
"""

import os
import re
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


# ─────────────────────────────────────────────
# Step 0.4: Text Cleaning & Normalization
# ─────────────────────────────────────────────

def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.strip()
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = text.replace('\u2018', "'").replace('\u2019', "'")
    text = text.replace('\u201c', '"').replace('\u201d', '"')
    text = text.replace('\u2013', '-').replace('\u2014', '--')
    text = text.replace('\u2026', '...')
    text = text.replace('\u2022', '-')
    text = text.replace('\u00a0', ' ')
    text = text.replace('\uf0b7', '-')
    text = text.replace('\u2010', '-').replace('\u2011', '-').replace('\u2012', '-')
    text = re.sub(r'[\u2000-\u200f\u2028-\u202f\u205f\u3000]', ' ', text)
    return text.strip()


# ─────────────────────────────────────────────
# Step 0.5: Document Metadata
# ─────────────────────────────────────────────

def get_doc_metadata(filepath: str):
    path = Path(filepath)
    fname = path.name
    mtime = datetime.fromtimestamp(path.stat().st_mtime).isoformat()

    doc_type_map = {
        "cheatsheet": "cheatsheet",
        "process flow": "process_flow",
        "training content": "training",
        "faq": "faq",
        "scenarios": "scenarios",
    }
    fname_lower = fname.lower()
    doc_type = "unknown"
    for key, val in doc_type_map.items():
        if key in fname_lower:
            doc_type = val
            break
    return {
        "source_doc": fname,
        "doc_type": doc_type,
        "last_modified": mtime,
        "version": "1.0",
    }


# ─────────────────────────────────────────────
# Step 0.1 & 0.2: Extract & Label Elements
# ─────────────────────────────────────────────

def detect_doc_type_from_content(paragraphs, tables) -> str:
    combined = " ".join(p.text.lower() for p in paragraphs[:50])
    if tables:
        return "glossary"
    if "scenario" in combined and ("what went wrong" in combined or "who was involved" in combined):
        return "scenarios"
    if "faq" in combined or re.search(r'(?m)^\d+\.\s+what (is|are)', combined):
        return "faq"
    if "slide" in combined and "narrator" in combined:
        return "process_flow"
    if "training" in combined or "module" in combined:
        return "training"
    return "unknown"


# ─────────────────────────────────────────────
# Step 0.3: Special Construct Handlers
# ─────────────────────────────────────────────

def extract_glossary_terms(table):
    rows = table.rows
    if len(rows) < 2:
        return []
    header_cells = [clean_text(cell.text).lower() for cell in rows[0].cells]
    col_map = {}
    for i, h in enumerate(header_cells):
        if "term" in h:
            col_map["term"] = i
        elif "full form" in h or "full" in h:
            col_map["full_form"] = i
        elif "explanation" in h or "explain" in h:
            col_map["explanation"] = i
        elif "example" in h:
            col_map["example"] = i
        elif "why" in h or "need" in h:
            col_map["why_needed"] = i
        elif "used by" in h or "used_by" in h:
            col_map["used_by"] = i

    if "term" not in col_map:
        return []

    terms = []
    for row in rows[1:]:
        cells = [clean_text(cell.text) for cell in row.cells]
        term = cells[col_map["term"]] if col_map["term"] < len(cells) else ""
        if not term:
            continue

        used_by_raw = cells[col_map.get("used_by", -1)] if col_map.get("used_by", -1) < len(cells) else ""
        used_by_list = [u.strip() for u in used_by_raw.split(",") if u.strip()] if used_by_raw else []

        entry = {
            "term": term,
            "full_form": cells[col_map.get("full_form", -1)] if col_map.get("full_form", -1) < len(cells) else "",
            "explanation": cells[col_map.get("explanation", -1)] if col_map.get("explanation", -1) < len(cells) else "",
            "example": cells[col_map.get("example", -1)] if col_map.get("example", -1) < len(cells) else "",
            "why_needed": cells[col_map.get("why_needed", -1)] if col_map.get("why_needed", -1) < len(cells) else "",
            "used_by": used_by_list,
        }
        terms.append(entry)
    return terms


def extract_scenarios(paragraphs):
    scenarios = []
    current = None
    FIELD_PATTERNS = [
        (r'^What Went Wrong\?\s*(.*)', "what_went_wrong"),
        (r'^Who Was Involved\?\s*(.*)', "who_was_involved"),
        (r'^Expected vs Actual:\s*(.*)', "expected_vs_actual"),
        (r'^Impact:\s*(.*)', "impact"),
        (r'^Comment Box \(Your Thoughts\):\s*(.*)', "comment_box"),
    ]

    def parse_scenario_header(txt: str):
        match = re.match(r'Scenario (\d+):\s*(.*)', txt, re.IGNORECASE)
        if not match:
            return None
        title = match.group(2).strip()
        remaining = txt[match.end():].strip()
        domain = ""
        scenario_text = ""
        if remaining:
            domain_match = re.search(r'Domain:\s*(.+?)(?:\n|$)', remaining, re.IGNORECASE)
            if domain_match:
                domain = clean_text(domain_match.group(1))
            scen_match = re.search(r'Scenario:\s*(.*)', remaining, re.DOTALL)
            if scen_match:
                scenario_text = clean_text(scen_match.group(1))
        return {
            "number": int(match.group(1)),
            "title": title,
            "domain": domain,
            "scenario_text": scenario_text,
        }

    def new_scenario(parsed):
        return {
            "scenario_number": parsed["number"],
            "title": parsed["title"],
            "domain": parsed["domain"],
            "scenario_text": parsed["scenario_text"],
            "what_went_wrong": "",
            "who_was_involved": "",
            "expected_vs_actual": "",
            "impact": "",
            "comment_box": "",
        }

    i = 0
    while i < len(paragraphs):
        txt = clean_text(paragraphs[i].text)
        if not txt:
            i += 1
            continue

        parsed = parse_scenario_header(txt)
        if parsed:
            if current:
                scenarios.append(current)
            # If domain/scenario_text not on the same line, peek at next paragraph
            if not parsed["domain"] and i + 1 < len(paragraphs):
                next_txt = clean_text(paragraphs[i + 1].text)
                next_parsed = parse_scenario_header(next_txt)
                if not next_parsed:
                    domain_match = re.search(r'Domain:\s*(.+?)(?:\n|$)', next_txt, re.IGNORECASE)
                    if domain_match:
                        parsed["domain"] = clean_text(domain_match.group(1))
                    scen_match = re.search(r'Scenario:\s*(.*)', next_txt, re.DOTALL)
                    if scen_match:
                        parsed["scenario_text"] = clean_text(scen_match.group(1))
                    # If the next paragraph was consumed, skip it
                    if domain_match or scen_match:
                        i += 1
            current = new_scenario(parsed)
            i += 1
            continue

        if current:
            for pat, field_name in FIELD_PATTERNS:
                m = re.match(pat, txt, re.DOTALL)
                if m:
                    current[field_name] = m.group(1).strip()
                    break

        i += 1

    if current:
        scenarios.append(current)
    return scenarios


def extract_terminology_from_text(text: str):
    match = re.match(r'Term(?:inology)?\s*(?:-|–|:)?\s*(.+?)\s*[:\-–]\s*(.*)', text, re.DOTALL)
    if match:
        return {
            "term": clean_text(match.group(1)),
            "definition": clean_text(match.group(2)),
        }
    match2 = re.match(r'Term(?:inology)?\s+(Explained)?:?\s*(.+?)\s*[:\-–]\s*(.*)', text, re.DOTALL)
    if match2:
        return {
            "term": clean_text(match2.group(2)),
            "definition": clean_text(match2.group(3)),
        }
    return None


def extract_qa(text: str, num: int):
    match = re.match(r'^(\d+)\.\s+(.*?)\?\s*(.*)', text, re.DOTALL)
    if match:
        q = clean_text(match.group(2) + "?")
        a = clean_text(match.group(3))
        return {"question_number": int(match.group(1)), "question": q, "answer": a}
    return None


# ─────────────────────────────────────────────
# Main Extraction Pipeline
# ─────────────────────────────────────────────

def process_document(filepath: str) -> list[dict]:
    doc = Document(filepath)
    metadata = get_doc_metadata(filepath)
    elements = []

    paragraphs = doc.paragraphs
    tables = doc.tables

    # Detect doc type from content
    detected_type = detect_doc_type_from_content(paragraphs, tables)
    if metadata["doc_type"] == "unknown":
        metadata["doc_type"] = detected_type

    # ── Handle glossary documents (tables) ──
    if metadata["doc_type"] == "cheatsheet" or (tables and metadata["doc_type"] == "glossary"):
        for p in paragraphs:
            txt = clean_text(p.text)
            if re.match(r'^\d+\.\s+', txt):
                elements.append({
                    "source_doc": metadata["source_doc"],
                    "doc_type": metadata["doc_type"],
                    "element_type": "section_heading",
                    "section_title": txt,
                    "text": txt,
                })

        for table in tables:
            terms = extract_glossary_terms(table)
            for term_entry in terms:
                elements.append({
                    "source_doc": metadata["source_doc"],
                    "doc_type": metadata["doc_type"],
                    "element_type": "term_definition",
                    **term_entry,
                })
        return elements

    # ── Handle scenarios documents ──
    if metadata["doc_type"] == "scenarios":
        scenarios = extract_scenarios(paragraphs)
        for s in scenarios:
            elements.append({
                "source_doc": metadata["source_doc"],
                "doc_type": metadata["doc_type"],
                "element_type": "scenario",
                **s,
            })
        return elements

    # ── Handle FAQ documents ──
    if metadata["doc_type"] == "faq":
        current_section = ""
        for p in paragraphs:
            txt = clean_text(p.text)
            if not txt:
                continue
            section_match = re.match(r'^([A-Z])\.\s+(.*)', txt)
            if section_match:
                current_section = f"{section_match.group(1)}. {section_match.group(2)}"
                elements.append({
                    "source_doc": metadata["source_doc"],
                    "doc_type": metadata["doc_type"],
                    "element_type": "section_heading",
                    "section_title": current_section,
                    "text": txt,
                })
                continue

            qa = extract_qa(txt, 0)
            if qa:
                elements.append({
                    "source_doc": metadata["source_doc"],
                    "doc_type": metadata["doc_type"],
                    "element_type": "qa_pair",
                    "section": current_section,
                    **qa,
                })
            else:
                elements.append({
                    "source_doc": metadata["source_doc"],
                    "doc_type": metadata["doc_type"],
                    "element_type": "paragraph",
                    "text": txt,
                })
        return elements

    # ── Handle process_flow documents ──
    if metadata["doc_type"] == "process_flow" or metadata["doc_type"] == "training":
        process_names = [
            "procurement", "inventory management", "order fulfillment",
            "sales and pos", "return management",
        ]
        current_process = ""
        current_slide = ""
        i = 0
        while i < len(paragraphs):
            txt = clean_text(paragraphs[i].text)
            if not txt:
                i += 1
                continue

            for pname in process_names:
                if pname in txt.lower() and ("process flow" in txt.lower() or "now let" in txt.lower()):
                    current_process = txt
                    elements.append({
                        "source_doc": metadata["source_doc"],
                        "doc_type": metadata["doc_type"],
                        "element_type": "process_section",
                        "process_name": txt,
                        "text": txt,
                    })
                    break

            slide_match = re.match(r'^Slide (\d+)[:\s-]*(.*)', txt, re.IGNORECASE)
            if slide_match:
                current_slide = f"Slide {slide_match.group(1)}"
                elements.append({
                    "source_doc": metadata["source_doc"],
                    "doc_type": metadata["doc_type"],
                    "element_type": "slide_heading",
                    "slide_number": int(slide_match.group(1)),
                    "slide_title": slide_match.group(2).strip() or current_slide,
                    "process_name": current_process,
                    "text": txt,
                })
                i += 1
                continue

            narrator_match = re.match(r'[🎤🎙️]\s*Narrator:\s*(.*)', txt, re.DOTALL)
            if narrator_match:
                elements.append({
                    "source_doc": metadata["source_doc"],
                    "doc_type": metadata["doc_type"],
                    "element_type": "narrator",
                    "narrator_text": clean_text(narrator_match.group(1)),
                    "slide": current_slide,
                    "process_name": current_process,
                })
                i += 1
                continue

            if txt.startswith("Actors:") or txt.startswith("Systems:"):
                elements.append({
                    "source_doc": metadata["source_doc"],
                    "doc_type": metadata["doc_type"],
                    "element_type": "section_label",
                    "label": txt.rstrip(":"),
                    "slide": current_slide,
                    "process_name": current_process,
                })
                i += 1
                continue

            term_def = extract_terminology_from_text(txt)
            if term_def:
                elements.append({
                    "source_doc": metadata["source_doc"],
                    "doc_type": metadata["doc_type"],
                    "element_type": "term_definition",
                    "process_name": current_process,
                    "slide": current_slide,
                    **term_def,
                })
                i += 1
                continue

            elements.append({
                "source_doc": metadata["source_doc"],
                "doc_type": metadata["doc_type"],
                "element_type": "paragraph",
                "process_name": current_process,
                "slide": current_slide,
                "text": txt,
            })
            i += 1

        return elements

    # ── Fallback: generic extraction ──
    for p in paragraphs:
        txt = clean_text(p.text)
        if not txt:
            continue
        elements.append({
            "source_doc": metadata["source_doc"],
            "doc_type": metadata["doc_type"],
            "element_type": "paragraph",
            "text": txt,
        })

    return elements


# ─────────────────────────────────────────────
# Step 0.6: Export to JSONL
# ─────────────────────────────────────────────

def export_jsonl(elements: list[dict], output_path: str):
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        for elem in elements:
            f.write(json.dumps(elem, ensure_ascii=False) + "\n")


# ─────────────────────────────────────────────
# CLI Entry Point
# ─────────────────────────────────────────────

def main():
    data_dir = Path(r"C:\Users\GanapathiSubramanian\Retail_agent\data")
    output_dir = Path(r"C:\Users\GanapathiSubramanian\Retail_agent\output\phase_0")
    summary_path = output_dir / "summary.json"

    docx_files = sorted(data_dir.glob("*.docx"))
    if not docx_files:
        print("No .docx files found in data directory.")
        return

    total_elements = 0
    file_counts = {}

    for fpath in docx_files:
        print(f"Processing: {fpath.name} ...")
        elements = process_document(str(fpath))
        base_name = fpath.stem.replace(" ", "_").lower()
        out_path = output_dir / f"{base_name}.jsonl"
        export_jsonl(elements, str(out_path))
        count = len(elements)
        file_counts[fpath.name] = count
        total_elements += count
        print(f"  -> {count} elements written to {out_path.name}")

    summary = {
        "total_files": len(docx_files),
        "total_elements": total_elements,
        "files": file_counts,
        "generated_at": datetime.now().isoformat(),
    }
    with open(summary_path, "w", encoding="utf-8") as f:
        json.dump(summary, f, indent=2, ensure_ascii=False)

    print(f"\n{'='*60}")
    print(f"Phase 0 Complete:")
    print(f"  Files processed: {len(docx_files)}")
    print(f"  Total elements: {total_elements}")
    print(f"  Output directory: {output_dir}")
    print(f"  Summary: {summary_path}")


if __name__ == "__main__":
    main()
